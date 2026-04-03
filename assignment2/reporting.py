from __future__ import annotations

import json
import textwrap
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches

from assignment2.config import DOC_DIR, FIGURES_DIR
from assignment2.helpers import clean_text


# Figure and report generation helpers.


def save_figures(clean_df: pd.DataFrame) -> list[Path]:
    sns.set_theme(style="whitegrid")
    figure_paths: list[Path] = []
    if clean_df.empty:
        return figure_paths

    charts = [
        (clean_df["source"].value_counts().rename_axis("source").reset_index(name="count"), "source", "count", "Vacancies by Source", FIGURES_DIR / "01_source_share.png"),
        (clean_df["city"].value_counts().head(10).rename_axis("city").reset_index(name="count"), "city", "count", "Top Cities", FIGURES_DIR / "02_top_cities.png"),
        (clean_df["role_category"].value_counts().head(8).rename_axis("role_category").reset_index(name="count"), "role_category", "count", "Role Categories", FIGURES_DIR / "03_role_categories.png"),
        (
            clean_df[clean_df["price_segment"].ne("")]["price_segment"].value_counts().rename_axis("price_segment").reset_index(name="count"),
            "price_segment",
            "count",
            "Price Segments",
            FIGURES_DIR / "04_price_segments.png",
        ),
    ]

    for chart_df, x_col, y_col, title, path in charts:
        if chart_df.empty:
            continue
        plt.figure(figsize=(10, 5))
        order = chart_df[x_col].tolist()
        sns.barplot(data=chart_df, x=x_col, y=y_col, hue=x_col, order=order, palette="viridis", legend=False)
        plt.xticks(rotation=25, ha="right")
        plt.title(title)
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        figure_paths.append(path)

    if clean_df["published_at"].notna().any():
        timeline = clean_df.assign(date=clean_df["published_at"].dt.date).groupby(["date", "source"]).size().reset_index(name="count")
        if not timeline.empty:
            path = FIGURES_DIR / "05_publication_timeline.png"
            plt.figure(figsize=(10, 5))
            sns.lineplot(data=timeline, x="date", y="count", hue="source", marker="o")
            plt.xticks(rotation=25, ha="right")
            plt.title("Vacancy Freshness Timeline")
            plt.tight_layout()
            plt.savefig(path, dpi=180)
            plt.close()
            figure_paths.append(path)

    coverage = clean_df.assign(has_salary=clean_df["salary_mid_kzt"].notna()).groupby(["source", "has_salary"]).size().reset_index(name="count")
    if not coverage.empty:
        path = FIGURES_DIR / "06_salary_transparency.png"
        plt.figure(figsize=(8, 5))
        sns.barplot(data=coverage, x="source", y="count", hue="has_salary", palette="mako")
        plt.title("Salary Transparency by Source")
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        figure_paths.append(path)

    return figure_paths


def build_platform_matrix() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "Platform": "hh.kz",
                "Student focus": "Medium",
                "Filters / search": "High",
                "Salary transparency": "Medium",
                "Technology level": "High",
                "Notes": "Strong search, commercial job board, but not niche for students.",
            },
            {
                "Platform": "Enbek.kz",
                "Student focus": "Medium-High",
                "Filters / search": "Medium",
                "Salary transparency": "High",
                "Technology level": "Medium",
                "Notes": "State platform with youth/practice routes, broad reach, weaker UX depth.",
            },
            {
                "Platform": "Staj.kz (target)",
                "Student focus": "Very High",
                "Filters / search": "Planned niche",
                "Salary transparency": "Very High",
                "Technology level": "Mobile-first niche",
                "Notes": "Niche platform around no-experience candidates, universities and feedback.",
            },
        ]
    )


def dataframe_to_markdown_table(df: pd.DataFrame) -> str:
    if df.empty:
        return "_No data_"
    columns = list(df.columns)
    lines = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for row in df.fillna("").astype(str).to_dict("records"):
        values = [row.get(column, "").replace("\n", " ").replace("|", "/") for column in columns]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def build_report_markdown(
    summary: dict[str, object],
    clean_df: pd.DataFrame,
    quality_df: pd.DataFrame,
    figure_paths: Iterable[Path],
) -> str:
    top_employers = clean_df["company"].value_counts().head(10).to_dict()
    top_skills = (
        clean_df["skills"].dropna().astype(str).str.split(",").explode().map(clean_text).replace("", pd.NA).dropna().value_counts().head(10).to_dict()
    )
    strengths = [
        "Итоговый датасет превышает требование задания по объему.",
        "Два источника дают более полную картину рынка, чем выборка только с одной платформы.",
        "Высокая доля вакансий без опыта подтверждает, что спрос на такой сегмент реально есть.",
    ]
    weaknesses = [
        "Часть вакансий публикуется без зарплаты, поэтому их труднее сравнивать между собой.",
        "Рынок ранней карьеры широкий, но не весь он состоит именно из стажировок.",
        "Крупные площадки не дают узкой настройки именно под студентов и выпускников.",
    ]
    opportunities = [
        "Сделать отдельный сервис для студентов с более прозрачной зарплатой и понятной фильтрацией.",
        "Начать со среднего ценового сегмента, где рынок выглядит наиболее устойчивым.",
        "Сфокусироваться сначала на Алматы и Астане, а потом расширяться дальше.",
    ]
    threats = [
        "Универсальные платформы уже собирают большую часть трафика вакансий.",
        "Рынок зависит от сезонности найма и от доли очных вакансий.",
        "Нишевому сервису нужно отдельно доказывать работодателям качество откликов.",
    ]
    figure_lines = "\n".join(f"- `{path.name}`" for path in figure_paths)
    return f"""# Assignment 2: Staj.kz Market Intelligence

## 1. Dataset summary
- Run date (UTC): `{summary['run_date_utc']}`
- Raw rows before cleaning: `{summary['raw_rows']}`
- Clean rows after cleaning: `{summary['clean_rows']}`
- Freshness window: `{summary['freshness_window_days']}` days
- Salary coverage: `{summary['salary_coverage_pct']}%`
- Student-friendly share: `{summary['student_friendly_pct']}%`
- Source split: `{json.dumps(summary['source_counts'], ensure_ascii=False)}`

## 2. Sources and scope
- Source 1: `hh.kz / api.hh.ru`
- Source 2: `Enbek.kz`
- Market scope: `early-career` vacancies for Kazakhstan
- Query logic: internship / trainee / graduate / junior / no-experience / youth / practice routes

## 3. Data quality log
{dataframe_to_markdown_table(quality_df)}

## 4. Market structure
- Top cities: `{json.dumps(summary['top_cities'], ensure_ascii=False)}`
- Top role categories: `{json.dumps(summary['top_role_categories'], ensure_ascii=False)}`
- Price segments: `{json.dumps(summary['price_segments'], ensure_ascii=False)}`
- Most promising segment for Staj.kz: `{summary['perspective_segment']}`

### Top employers
`{json.dumps(top_employers, ensure_ascii=False)}`

### Top extracted skills
`{json.dumps(top_skills, ensure_ascii=False)}`

## 5. Competitor analysis
### Platform matrix
{dataframe_to_markdown_table(build_platform_matrix())}

### SWOT
**Strengths**
{chr(10).join(f"- {item}" for item in strengths)}

**Weaknesses**
{chr(10).join(f"- {item}" for item in weaknesses)}

**Opportunities**
{chr(10).join(f"- {item}" for item in opportunities)}

**Threats**
{chr(10).join(f"- {item}" for item in threats)}

## 6. Strategic recommendation
The analysis suggests starting Staj.kz in the `{summary['perspective_segment']}` segment. It gives enough market volume for launch and looks more realistic for an initial market entry than the narrow premium tier. The first priority cities should be the largest hiring hubs, especially Almaty and Astana.

## 7. Figures
{figure_lines}
"""


def build_report_docx(markdown_report: str, figure_paths: Iterable[Path], output_path: Path) -> None:
    document = Document()
    document.add_heading("Assignment 2. Staj.kz Market Intelligence", level=0)
    for paragraph in markdown_report.split("\n\n"):
        text = paragraph.strip()
        if not text:
            continue
        if text.startswith("# "):
            document.add_heading(text[2:].strip(), level=1)
        elif text.startswith("## "):
            document.add_heading(text[3:].strip(), level=2)
        elif text.startswith("### "):
            document.add_heading(text[4:].strip(), level=3)
        else:
            lines = text.splitlines()
            if all(line.startswith("- ") for line in lines):
                for line in lines:
                    document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(textwrap.dedent(text))

    document.add_heading("Figures", level=2)
    for figure in figure_paths:
        if figure.exists():
            document.add_paragraph(figure.name)
            document.add_picture(str(figure), width=Inches(6.4))

    document.save(output_path)
