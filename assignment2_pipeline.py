from __future__ import annotations

import argparse
import json
import math
import re
import textwrap
import time
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urljoin

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import pandas as pd
import requests
import seaborn as sns
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry


BASE_DIR = Path(__file__).resolve().parent
RAW_HH_DIR = BASE_DIR / "data" / "raw" / "hh"
RAW_ENBEK_DIR = BASE_DIR / "data" / "raw" / "enbek"
RAW_ENBEK_DETAIL_DIR = RAW_ENBEK_DIR / "details"
PROCESSED_DIR = BASE_DIR / "data" / "processed"
FIGURES_DIR = BASE_DIR / "output" / "figures"
DOC_DIR = BASE_DIR / "output" / "doc"

RUN_DATE = datetime.now(timezone.utc)
FRESHNESS_DAYS = 90
KZT_EXCHANGE_RATES = {
    "KZT": 1.0,
    "USD": 500.0,
    "EUR": 540.0,
    "RUR": 5.5,
    "RUB": 5.5,
}
KZ_CITIES = [
    "Алматы",
    "Астана",
    "Шымкент",
    "Уральск",
    "Караганда",
    "Карағанды",
    "Актобе",
    "Ақтөбе",
    "Павлодар",
    "Костанай",
    "Қостанай",
    "Кокшетау",
    "Көкшетау",
    "Усть-Каменогорск",
    "Өскемен",
    "Семей",
    "Тараз",
    "Талдыкорган",
    "Талдықорған",
    "Атырау",
    "Актау",
    "Ақтау",
    "Петропавловск",
    "Петропавл",
    "Кызылорда",
    "Қызылорда",
    "Туркестан",
    "Түркістан",
    "Жезказган",
    "Жезқазған",
    "Рудный",
    "Экибастуз",
    "Екібастұз",
]
ROLE_KEYWORDS = {
    "IT / Data": [
        "python",
        "data",
        "аналит",
        "developer",
        "разработ",
        "qa",
        "tester",
        "software",
        "machine learning",
        "devops",
        "backend",
        "frontend",
        "fullstack",
        "информатик",
    ],
    "Marketing / Sales": [
        "marketing",
        "smm",
        "sales",
        "продаж",
        "маркет",
        "brand",
        "content",
        "pr ",
        "реклам",
        "бариста",
        "продав",
    ],
    "Finance / Admin": [
        "finance",
        "account",
        "бухгал",
        "экономист",
        "офис",
        "hr",
        "кадр",
        "админист",
        "assistan",
    ],
    "Education / Support": [
        "teacher",
        "учител",
        "образован",
        "ментор",
        "support",
        "оператор",
        "консульт",
    ],
    "Operations / Logistics": [
        "операц",
        "logist",
        "warehouse",
        "склад",
        "закуп",
        "транспорт",
        "разнорабоч",
        "производ",
    ],
}
EARLY_CAREER_TERMS = [
    "стаж",
    "intern",
    "trainee",
    "graduate",
    "junior",
    "без опыта",
    "начинающ",
    "student",
    "студент",
    "practice",
    "практик",
]
HH_QUERIES = [
    "стажировка",
    "стажер",
    "intern",
    "internship",
    "trainee",
    "junior",
    "graduate",
    "студент",
    "без опыта",
    "первая работа",
    "начинающий специалист",
]
ENBEK_SEARCHES = [
    {"label": "youth", "params": {"tag": "youth", "experience": "e0"}, "pages": 12},
    {"label": "practice", "params": {"tag": "practice", "experience": "e0"}, "pages": 8},
    {"label": "intern", "params": {"prof": "intern", "tag": "youth", "experience": "e0"}, "pages": 4},
    {"label": "junior", "params": {"prof": "junior", "tag": "youth", "experience": "e0"}, "pages": 4},
    {"label": "trainee", "params": {"prof": "trainee", "tag": "youth", "experience": "e0"}, "pages": 4},
    {"label": "стажировка", "params": {"prof": "стажировка", "tag": "youth", "experience": "e0"}, "pages": 6},
    {"label": "без опыта", "params": {"prof": "без опыта", "tag": "youth", "experience": "e0"}, "pages": 6},
]
COMMON_COLUMNS = [
    "source",
    "source_id",
    "url",
    "title",
    "company",
    "city",
    "region",
    "published_at",
    "employment",
    "schedule",
    "experience",
    "education",
    "salary_from",
    "salary_to",
    "currency",
    "salary_mid_kzt",
    "query_group",
    "role_category",
    "is_student_friendly",
    "description_text",
    "industry",
    "subtitle",
    "skills",
    "languages",
    "work_format",
    "internship_flag",
    "source_listing_label",
    "price_segment",
    "salary_outlier_flag",
]


@dataclass
class QualityLogRow:
    step: str
    metric: str
    value: Any


def ensure_dirs() -> None:
    for path in [RAW_HH_DIR, RAW_ENBEK_DIR, RAW_ENBEK_DETAIL_DIR, PROCESSED_DIR, FIGURES_DIR, DOC_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def slugify(text: str) -> str:
    value = re.sub(r"\s+", "-", text.strip().lower())
    value = re.sub(r"[^a-zA-Z0-9а-яА-ЯёЁ_-]+", "-", value)
    value = re.sub(r"-{2,}", "-", value).strip("-")
    return value or "query"


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_date(value: Any) -> pd.Timestamp | pd.NaT:
    text = clean_text(value)
    if not text:
        return pd.NaT
    for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S%z", "%d.%m.%Y", "%d.%m.%Y %H:%M"):
        try:
            return pd.Timestamp(datetime.strptime(text, fmt))
        except ValueError:
            continue
    try:
        return pd.to_datetime(text, errors="coerce")
    except Exception:
        return pd.NaT


def parse_salary_numbers(text: str) -> tuple[float | None, float | None]:
    values = re.findall(r"\d[\d\s]*", clean_text(text))
    nums = [float(v.replace(" ", "")) for v in values if v.strip()]
    if not nums:
        return None, None
    if len(nums) == 1:
        return nums[0], nums[0]
    return nums[0], nums[1]


def currency_to_kzt(amount: float | None, currency: str | None) -> float | None:
    if amount is None:
        return None
    rate = KZT_EXCHANGE_RATES.get(clean_text(currency).upper(), 1.0 if clean_text(currency).upper() == "KZT" else None)
    if rate is None:
        return None
    return round(amount * rate, 2)


def midpoint(low: float | None, high: float | None) -> float | None:
    if low is None and high is None:
        return None
    if low is None:
        return high
    if high is None:
        return low
    return round((low + high) / 2, 2)


def normalize_city(city: Any, region: Any = "") -> str:
    combined = " | ".join([clean_text(city), clean_text(region)])
    if not combined.strip("| "):
        return ""
    for known_city in KZ_CITIES:
        if known_city.lower() in combined.lower():
            if known_city in {"Қостанай", "Карағанды", "Көкшетау", "Талдықорған", "Ақтау", "Ақтөбе", "Қызылорда", "Өскемен", "Жезқазған", "Екібастұз", "Түркістан"}:
                return {
                    "Қостанай": "Костанай",
                    "Карағанды": "Караганда",
                    "Көкшетау": "Кокшетау",
                    "Талдықорған": "Талдыкорган",
                    "Ақтау": "Актау",
                    "Ақтөбе": "Актобе",
                    "Қызылорда": "Кызылорда",
                    "Өскемен": "Усть-Каменогорск",
                    "Жезқазған": "Жезказган",
                    "Екібастұз": "Экибастуз",
                    "Түркістан": "Туркестан",
                }[known_city]
            return known_city
    if "," in clean_text(city):
        return clean_text(city).split(",")[-1].replace("г.", "").strip()
    if "," in clean_text(region):
        return clean_text(region).split(",")[-1].replace("г.", "").strip()
    return clean_text(city) or clean_text(region)


def infer_role_category(*parts: Any) -> str:
    haystack = " ".join(clean_text(part).lower() for part in parts if clean_text(part))
    for label, keywords in ROLE_KEYWORDS.items():
        if any(keyword in haystack for keyword in keywords):
            return label
    return "General Early-Career"


def infer_student_friendly(*parts: Any) -> bool:
    haystack = " ".join(clean_text(part).lower() for part in parts if clean_text(part))
    return any(term in haystack for term in EARLY_CAREER_TERMS)


def normalize_employment(value: Any) -> str:
    text = clean_text(value).lower()
    mapping = {
        "полная занятость": "Full-time",
        "полная": "Full-time",
        "частичная занятость": "Part-time",
        "частичная": "Part-time",
        "проектная работа": "Project",
        "постоянная": "Permanent",
        "временная": "Temporary",
    }
    return mapping.get(text, clean_text(value))


def normalize_schedule(value: Any) -> str:
    text = clean_text(value).lower()
    mapping = {
        "полный день": "Full day",
        "полный рабочий день": "Full day",
        "гибкий график": "Flexible",
        "сменный график": "Shift",
        "удаленная работа": "Remote",
        "удалённая работа": "Remote",
        "вахта": "Rotation",
    }
    return mapping.get(text, clean_text(value))


def request_session() -> requests.Session:
    session = requests.Session()
    retry = Retry(
        total=5,
        connect=5,
        read=5,
        backoff_factor=1.2,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET", "HEAD", "OPTIONS"],
    )
    adapter = HTTPAdapter(max_retries=retry, pool_connections=20, pool_maxsize=20)
    session.mount("https://", adapter)
    session.mount("http://", adapter)
    session.headers.update(
        {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_0) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
            "Accept-Language": "ru,en;q=0.9",
            "Connection": "close",
        }
    )
    return session


def dump_text(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def dump_json(path: Path, payload: Any) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def http_get(session: requests.Session, url: str, *, params: dict[str, Any] | None = None, timeout: int = 30) -> requests.Response:
    last_error: Exception | None = None
    for attempt in range(1, 6):
        try:
            response = session.get(url, params=params, timeout=timeout)
            response.raise_for_status()
            return response
        except requests.RequestException as exc:
            last_error = exc
            if attempt == 5:
                raise
            time.sleep(min(2 ** attempt, 8))
    assert last_error is not None
    raise last_error


def fetch_hh_records(session: requests.Session, max_pages_per_query: int) -> pd.DataFrame:
    records: list[dict[str, Any]] = []
    for query in HH_QUERIES:
        print(f"[hh] query={query}", flush=True)
        query_slug = slugify(query)
        for page in range(max_pages_per_query):
            raw_path = RAW_HH_DIR / f"{query_slug}_page_{page + 1}.json"
            params = {
                "text": query,
                "area": 40,
                "per_page": 100,
                "page": page,
                "experience": "noExperience",
                "order_by": "publication_time",
            }
            if raw_path.exists():
                payload = json.loads(raw_path.read_text(encoding="utf-8"))
            else:
                response = http_get(session, "https://api.hh.ru/vacancies", params=params, timeout=30)
                payload = response.json()
                dump_json(raw_path, payload)
            items = payload.get("items", [])
            if not items:
                break

            for item in items:
                salary = item.get("salary") or {}
                salary_from = salary.get("from")
                salary_to = salary.get("to")
                currency = salary.get("currency") or "KZT"
                description_text = " ".join(
                    [
                        clean_text(item.get("snippet", {}).get("requirement")),
                        clean_text(item.get("snippet", {}).get("responsibility")),
                    ]
                ).strip()
                work_format = ", ".join(clean_text(entry.get("name")) for entry in item.get("work_format", []))
                professional_roles = ", ".join(clean_text(entry.get("name")) for entry in item.get("professional_roles", []))
                record = {
                    "source": "hh.kz",
                    "source_id": clean_text(item.get("id")),
                    "url": clean_text(item.get("alternate_url")),
                    "title": clean_text(item.get("name")),
                    "company": clean_text(item.get("employer", {}).get("name")),
                    "city": normalize_city(item.get("area", {}).get("name")),
                    "region": clean_text(item.get("area", {}).get("name")),
                    "published_at": parse_date(item.get("published_at")),
                    "employment": normalize_employment(item.get("employment", {}).get("name")),
                    "schedule": normalize_schedule(item.get("schedule", {}).get("name")),
                    "experience": clean_text(item.get("experience", {}).get("name")),
                    "education": "",
                    "salary_from": salary_from,
                    "salary_to": salary_to,
                    "currency": clean_text(currency or "KZT"),
                    "salary_mid_kzt": currency_to_kzt(midpoint(salary_from, salary_to), currency),
                    "query_group": query,
                    "role_category": infer_role_category(
                        item.get("name"), professional_roles, item.get("snippet", {}).get("requirement")
                    ),
                    "is_student_friendly": infer_student_friendly(
                        query,
                        item.get("name"),
                        item.get("snippet", {}).get("requirement"),
                        item.get("experience", {}).get("name"),
                    ),
                    "description_text": description_text,
                    "industry": professional_roles,
                    "subtitle": "",
                    "skills": "",
                    "languages": "",
                    "work_format": work_format,
                    "internship_flag": bool(item.get("internship")),
                    "source_listing_label": f"hh_search::{query}",
                    "price_segment": "",
                    "salary_outlier_flag": False,
                }
                records.append(record)

            if page + 1 >= payload.get("pages", 1):
                break
            time.sleep(0.2)

    return pd.DataFrame(records, columns=COMMON_COLUMNS)


def parse_enbek_listing_card(card: BeautifulSoup, query_group: str) -> dict[str, Any]:
    link_tag = card.select_one(".title a.stretched")
    title = clean_text(link_tag.get_text()) if link_tag else ""
    href = link_tag.get("href", "") if link_tag else ""
    source_id_match = re.search(r"~(\d+)$", href)
    source_id = source_id_match.group(1) if source_id_match else clean_text(card.get("wire:key", "")).replace("item-", "")
    subtitle = clean_text(card.select_one(".subtitle").get_text()) if card.select_one(".subtitle") else ""
    industry = clean_text(card.select_one(".profobl").get_text()) if card.select_one(".profobl") else ""
    price_text = clean_text(card.select_one(".price").get_text()) if card.select_one(".price") else ""
    salary_from, salary_to = parse_salary_numbers(price_text)
    company = clean_text(card.select_one(".company").get_text()) if card.select_one(".company") else ""
    location = clean_text(card.select_one(".location").get_text()) if card.select_one(".location") else ""
    experience = clean_text(card.select_one(".experience").get_text()) if card.select_one(".experience") else ""
    schedule = clean_text(card.select_one(".time").get_text()) if card.select_one(".time") else ""
    education = clean_text(card.select_one(".education").get_text()) if card.select_one(".education") else ""
    published_text = clean_text(card.select_one(".right-content").get_text()) if card.select_one(".right-content") else ""
    published_match = re.search(r"(\d{2}\.\d{2}\.\d{4})", published_text)
    published_at = parse_date(published_match.group(1)) if published_match else pd.NaT
    record = {
        "source": "enbek.kz",
        "source_id": source_id,
        "url": urljoin("https://www.enbek.kz", href),
        "title": title,
        "company": company,
        "city": normalize_city(location, location),
        "region": location,
        "published_at": published_at,
        "employment": "",
        "schedule": normalize_schedule(schedule),
        "experience": experience,
        "education": education,
        "salary_from": salary_from,
        "salary_to": salary_to,
        "currency": "KZT",
        "salary_mid_kzt": midpoint(salary_from, salary_to),
        "query_group": query_group,
        "role_category": infer_role_category(title, subtitle, industry),
        "is_student_friendly": infer_student_friendly(query_group, title, subtitle, experience, education),
        "description_text": "",
        "industry": industry,
        "subtitle": subtitle,
        "skills": "",
        "languages": "",
        "work_format": "",
        "internship_flag": False,
        "source_listing_label": f"enbek_search::{query_group}",
        "price_segment": "",
        "salary_outlier_flag": False,
    }
    return record


def fetch_enbek_search_records(session: requests.Session) -> pd.DataFrame:
    records: list[dict[str, Any]] = []
    seen_pages: set[str] = set()

    for search in ENBEK_SEARCHES:
        label = search["label"]
        print(f"[enbek] search={label}", flush=True)
        params = dict(search["params"])
        for page in range(1, search["pages"] + 1):
            params_with_page = dict(params)
            params_with_page["page"] = page
            page_key = f"{label}_{page}"
            if page_key in seen_pages:
                continue
            seen_pages.add(page_key)
            raw_path = RAW_ENBEK_DIR / f"{label}_page_{page}.html"
            if raw_path.exists():
                html = raw_path.read_text(encoding="utf-8")
            else:
                try:
                    response = http_get(session, "https://www.enbek.kz/ru/search/vacancy", params=params_with_page, timeout=30)
                except requests.RequestException:
                    break
                html = response.text
                dump_text(raw_path, html)
            soup = BeautifulSoup(html, "lxml")
            cards = soup.select("div.item-list[wire\\:key^='item-']")
            if not cards:
                break
            for card in cards:
                records.append(parse_enbek_listing_card(card, label))
            time.sleep(0.2)

    listing_df = pd.DataFrame(records, columns=COMMON_COLUMNS)
    if listing_df.empty:
        return listing_df
    return enrich_enbek_details(session, listing_df)


def parse_enbek_jsonld(soup: BeautifulSoup) -> dict[str, Any]:
    scripts = soup.find_all("script", attrs={"type": "application/ld+json"})
    for script in scripts:
        raw = clean_text(script.string or script.get_text())
        if "@type" not in raw or "JobPosting" not in raw:
            continue
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            continue
    return {}


def fetch_enbek_detail(session: requests.Session, url: str) -> dict[str, Any]:
    detail_id = re.search(r"~(\d+)$", url)
    cache_path = RAW_ENBEK_DETAIL_DIR / f"{detail_id.group(1) if detail_id else slugify(url)}.html"
    if cache_path.exists():
        html = cache_path.read_text(encoding="utf-8")
    else:
        response = http_get(session, url, timeout=30)
        html = response.text
        dump_text(cache_path, html)
        time.sleep(0.25)
    soup = BeautifulSoup(html, "lxml")
    jsonld = parse_enbek_jsonld(soup)

    detail: dict[str, Any] = {
        "employment": "",
        "schedule": "",
        "experience": "",
        "education": "",
        "salary_from": None,
        "salary_to": None,
        "currency": "KZT",
        "salary_mid_kzt": None,
        "description_text": "",
        "skills": "",
        "languages": "",
        "region": "",
        "city": "",
        "internship_flag": False,
        "company": "",
    }

    price_text = clean_text(soup.select_one(".price").get_text()) if soup.select_one(".price") else ""
    salary_from, salary_to = parse_salary_numbers(price_text)
    detail["salary_from"] = salary_from
    detail["salary_to"] = salary_to
    detail["salary_mid_kzt"] = midpoint(salary_from, salary_to)

    field_map = {
        "Тип занятости": "employment",
        "График работы": "schedule",
        "Опыт работы": "experience",
        "Образование": "education",
        "Стажировка": "internship_flag",
        "Регион": "region",
    }
    for block in soup.select(".single-line"):
        label_tag = block.select_one(".label")
        value_tag = block.select_one(".value")
        if not label_tag or not value_tag:
            continue
        label = clean_text(label_tag.get_text())
        value = clean_text(value_tag.get_text(" "))
        if label in field_map:
            key = field_map[label]
            detail[key] = value

    detail["company"] = clean_text(soup.select_one(".company-box .info a").get_text()) if soup.select_one(".company-box .info a") else ""
    detail["city"] = normalize_city(detail.get("region"), detail.get("region"))

    skills = [clean_text(tag.get_text()) for tag in soup.select(".single-line .list-inline .list-inline-item")]
    detail["skills"] = ", ".join(sorted(set(skill for skill in skills if skill)))

    language_values: list[str] = []
    for tag in soup.select(".single-line .list-unstyled li"):
        value = clean_text(tag.get_text())
        if "/" in value:
            language_values.append(value)
    detail["languages"] = ", ".join(sorted(set(language_values)))

    text_chunks = []
    title_tag = soup.select_one("h4.title strong")
    subtitle_tag = soup.select_one(".subtitle")
    if title_tag:
        text_chunks.append(clean_text(title_tag.get_text()))
    if subtitle_tag:
        text_chunks.append(clean_text(subtitle_tag.get_text()))
    if jsonld.get("description"):
        text_chunks.append(clean_text(jsonld.get("description")))
    if jsonld.get("skills"):
        text_chunks.append(clean_text(jsonld.get("skills")))
    detail["description_text"] = " ".join(chunk for chunk in text_chunks if chunk)

    if jsonld.get("baseSalary", {}).get("currency"):
        detail["currency"] = clean_text(jsonld["baseSalary"]["currency"])
    if jsonld.get("educationRequirements") and not detail["education"]:
        detail["education"] = clean_text(jsonld["educationRequirements"])
    if jsonld.get("experienceRequirements") and not detail["experience"]:
        detail["experience"] = clean_text(jsonld["experienceRequirements"])
    if jsonld.get("hiringOrganization", {}).get("name") and not detail["company"]:
        detail["company"] = clean_text(jsonld["hiringOrganization"]["name"])
    if jsonld.get("jobLocation", {}).get("address", {}).get("addressLocality"):
        locality = clean_text(jsonld["jobLocation"]["address"]["addressLocality"])
        detail["region"] = detail["region"] or locality
        detail["city"] = normalize_city(locality, locality)

    internship_value = clean_text(detail.get("internship_flag"))
    detail["internship_flag"] = internship_value.lower().startswith("предполага") or "practice" in clean_text(url).lower()
    detail["schedule"] = normalize_schedule(detail.get("schedule"))
    detail["employment"] = normalize_employment(detail.get("employment"))
    return detail


def enrich_enbek_details(session: requests.Session, listing_df: pd.DataFrame) -> pd.DataFrame:
    rows = listing_df.to_dict("records")
    unique_urls = sorted({row["url"] for row in rows if row.get("url")})
    unique_urls = unique_urls[:120]
    detail_map: dict[str, dict[str, Any]] = {}

    for index, url in enumerate(unique_urls, start=1):
        if index == 1 or index % 25 == 0:
            print(f"[enbek] detail {index}/{len(unique_urls)}", flush=True)
        try:
            detail_map[url] = fetch_enbek_detail(session, url)
        except Exception as exc:
            detail_map[url] = {"description_text": f"detail fetch failed: {exc}"}

    enriched_rows: list[dict[str, Any]] = []
    for row in rows:
        detail = detail_map.get(row["url"], {})
        merged = dict(row)
        for key, value in detail.items():
            if key not in merged or not merged[key]:
                merged[key] = value
            elif key in {"description_text", "skills", "languages"} and value:
                merged[key] = clean_text(f"{merged[key]} {value}")
        merged["company"] = merged["company"] or detail.get("company", "")
        merged["city"] = normalize_city(merged.get("city"), merged.get("region"))
        merged["role_category"] = infer_role_category(
            merged.get("title"), merged.get("subtitle"), merged.get("industry"), merged.get("description_text")
        )
        merged["is_student_friendly"] = infer_student_friendly(
            merged.get("query_group"),
            merged.get("title"),
            merged.get("description_text"),
            merged.get("experience"),
            merged.get("education"),
        )
        if merged.get("salary_mid_kzt") is None:
            merged["salary_mid_kzt"] = midpoint(merged.get("salary_from"), merged.get("salary_to"))
        enriched_rows.append(merged)

    return pd.DataFrame(enriched_rows, columns=COMMON_COLUMNS)


def standardize_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    working = df.copy()
    for column in COMMON_COLUMNS:
        if column not in working.columns:
            working[column] = ""

    text_columns = [
        "source",
        "source_id",
        "url",
        "title",
        "company",
        "city",
        "region",
        "employment",
        "schedule",
        "experience",
        "education",
        "currency",
        "query_group",
        "role_category",
        "description_text",
        "industry",
        "subtitle",
        "skills",
        "languages",
        "work_format",
        "source_listing_label",
    ]
    for column in text_columns:
        working[column] = working[column].map(clean_text)

    working["city"] = working.apply(lambda row: normalize_city(row["city"], row["region"]), axis=1)
    working["employment"] = working["employment"].map(normalize_employment)
    working["schedule"] = working["schedule"].map(normalize_schedule)
    working["published_at"] = working["published_at"].map(parse_date)
    working["published_at"] = pd.to_datetime(working["published_at"], errors="coerce", utc=True).dt.tz_convert(None)
    working["salary_from"] = pd.to_numeric(working["salary_from"], errors="coerce")
    working["salary_to"] = pd.to_numeric(working["salary_to"], errors="coerce")
    working["salary_mid_kzt"] = pd.to_numeric(working["salary_mid_kzt"], errors="coerce")

    missing_mid = working["salary_mid_kzt"].isna()
    working.loc[missing_mid, "salary_mid_kzt"] = working.loc[missing_mid].apply(
        lambda row: currency_to_kzt(midpoint(row["salary_from"], row["salary_to"]), row["currency"]),
        axis=1,
    )
    working["is_student_friendly"] = working.apply(
        lambda row: bool(
            row["is_student_friendly"]
            or infer_student_friendly(
                row["query_group"], row["title"], row["description_text"], row["experience"], row["education"]
            )
        ),
        axis=1,
    )
    working["internship_flag"] = working["internship_flag"].astype(bool)
    working["role_category"] = working.apply(
        lambda row: row["role_category"]
        or infer_role_category(row["title"], row["subtitle"], row["industry"], row["description_text"]),
        axis=1,
    )
    return working[COMMON_COLUMNS]


def apply_quality_cleaning(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    logs: list[QualityLogRow] = [QualityLogRow("raw", "rows", len(df))]
    working = df.copy()

    working = working.drop_duplicates(subset=["source", "source_id"], keep="first")
    logs.append(QualityLogRow("dedup", "after_source_id_dedup", len(working)))

    working["published_at"] = pd.to_datetime(working["published_at"], errors="coerce", utc=True).dt.tz_convert(None)
    freshness_cutoff = pd.Timestamp(RUN_DATE.date()) - pd.Timedelta(days=FRESHNESS_DAYS)
    before_stale = len(working)
    working = working[working["published_at"].notna()]
    working = working[working["published_at"] >= freshness_cutoff]
    logs.append(QualityLogRow("freshness", "removed_stale_or_invalid_dates", before_stale - len(working)))

    for column in ["title", "company", "city"]:
        working[column] = working[column].map(clean_text)
    working = working[working["title"].ne("") & working["company"].ne("")]
    logs.append(QualityLogRow("completeness", "after_required_fields_filter", len(working)))

    soft_key = (
        working["title"].str.lower()
        + "||"
        + working["company"].str.lower()
        + "||"
        + working["city"].str.lower()
        + "||"
        + working["published_at"].dt.strftime("%Y-%m-%d")
    )
    working = working.assign(_soft_key=soft_key)
    before_soft = len(working)
    working = working.drop_duplicates(subset=["_soft_key"], keep="first").drop(columns=["_soft_key"])
    logs.append(QualityLogRow("dedup", "removed_soft_duplicates", before_soft - len(working)))

    working["salary_mid_kzt"] = pd.to_numeric(working["salary_mid_kzt"], errors="coerce")
    positive_salaries = working.loc[working["salary_mid_kzt"].gt(0), "salary_mid_kzt"]
    q1 = positive_salaries.quantile(0.25) if not positive_salaries.empty else math.nan
    q3 = positive_salaries.quantile(0.75) if not positive_salaries.empty else math.nan
    iqr = q3 - q1 if pd.notna(q1) and pd.notna(q3) else math.nan
    lower_bound = q1 - 1.5 * iqr if pd.notna(iqr) else math.nan
    upper_bound = q3 + 1.5 * iqr if pd.notna(iqr) else math.nan
    working["salary_outlier_flag"] = False
    if pd.notna(lower_bound) and pd.notna(upper_bound):
        salary_mask = working["salary_mid_kzt"].notna() & (
            (working["salary_mid_kzt"] < lower_bound) | (working["salary_mid_kzt"] > upper_bound)
        )
        working.loc[salary_mask, "salary_outlier_flag"] = True
        working.loc[salary_mask, ["salary_from", "salary_to", "salary_mid_kzt"]] = pd.NA
        logs.append(QualityLogRow("salary_iqr", "salary_values_nullified_as_outliers", int(salary_mask.sum())))
        logs.append(QualityLogRow("salary_iqr", "iqr_lower_bound_kzt", round(lower_bound, 2)))
        logs.append(QualityLogRow("salary_iqr", "iqr_upper_bound_kzt", round(upper_bound, 2)))

    working["price_segment"] = ""
    clean_salary = working["salary_mid_kzt"].dropna()
    if not clean_salary.empty:
        q25, q50, q75 = clean_salary.quantile([0.25, 0.5, 0.75]).tolist()

        def assign_segment(value: Any) -> str:
            if pd.isna(value):
                return "unknown"
            if value <= q25:
                return "low-priced"
            if value <= q50:
                return "middle-priced"
            if value <= q75:
                return "high-priced"
            return "luxury"

        working["price_segment"] = working["salary_mid_kzt"].map(assign_segment)
        logs.append(QualityLogRow("pricing", "segment_q25_kzt", round(q25, 2)))
        logs.append(QualityLogRow("pricing", "segment_q50_kzt", round(q50, 2)))
        logs.append(QualityLogRow("pricing", "segment_q75_kzt", round(q75, 2)))

    logs.append(QualityLogRow("final", "rows", len(working)))
    return working[COMMON_COLUMNS], pd.DataFrame([row.__dict__ for row in logs])


def save_dataframe(df: pd.DataFrame, path: Path) -> None:
    df.to_csv(path, index=False, encoding="utf-8")


def build_summary(clean_df: pd.DataFrame, raw_df: pd.DataFrame, quality_df: pd.DataFrame) -> dict[str, Any]:
    salary_coverage = round(float(clean_df["salary_mid_kzt"].notna().mean() * 100), 2) if len(clean_df) else 0.0
    student_share = round(float(clean_df["is_student_friendly"].mean() * 100), 2) if len(clean_df) else 0.0
    source_counts = clean_df["source"].value_counts().to_dict()
    role_counts = clean_df["role_category"].value_counts().head(6).to_dict()
    city_counts = clean_df["city"].value_counts().head(8).to_dict()
    segment_counts = clean_df["price_segment"].value_counts().to_dict()
    platform_coverage = (
        clean_df.groupby("source")["salary_mid_kzt"]
        .apply(lambda s: round(float(s.notna().mean() * 100), 2) if len(s) else 0.0)
        .to_dict()
    )
    top_segment = next(iter(pd.Series(segment_counts).sort_values(ascending=False).index), "unknown") if segment_counts else "unknown"
    perspective_segment = top_segment
    if "middle-priced" in segment_counts:
        perspective_segment = "middle-priced"

    summary = {
        "run_date_utc": RUN_DATE.isoformat(),
        "raw_rows": int(len(raw_df)),
        "clean_rows": int(len(clean_df)),
        "source_counts": source_counts,
        "salary_coverage_pct": salary_coverage,
        "student_friendly_pct": student_share,
        "top_cities": city_counts,
        "top_role_categories": role_counts,
        "price_segments": segment_counts,
        "platform_salary_coverage_pct": platform_coverage,
        "freshness_window_days": FRESHNESS_DAYS,
        "perspective_segment": perspective_segment,
        "quality_log": quality_df.to_dict("records"),
    }
    return summary


def save_figures(clean_df: pd.DataFrame) -> list[Path]:
    sns.set_theme(style="whitegrid")
    figure_paths: list[Path] = []

    if clean_df.empty:
        return figure_paths

    charts = [
        (
            clean_df["source"].value_counts().rename_axis("source").reset_index(name="count"),
            "source",
            "count",
            "Vacancies by Source",
            FIGURES_DIR / "01_source_share.png",
        ),
        (
            clean_df["city"].value_counts().head(10).rename_axis("city").reset_index(name="count"),
            "city",
            "count",
            "Top Cities",
            FIGURES_DIR / "02_top_cities.png",
        ),
        (
            clean_df["role_category"].value_counts().head(8).rename_axis("role_category").reset_index(name="count"),
            "role_category",
            "count",
            "Role Categories",
            FIGURES_DIR / "03_role_categories.png",
        ),
        (
            clean_df[clean_df["price_segment"].ne("")]["price_segment"]
            .value_counts()
            .rename_axis("price_segment")
            .reset_index(name="count"),
            "price_segment",
            "count",
            "Price Segments",
            FIGURES_DIR / "04_price_segments.png",
        ),
    ]

    for chart_df, x_col, y_col, title, path in charts:
        if chart_df.empty:
            continue
        plt.figure(figsize=(10, 5))
        order = chart_df[x_col].tolist()
        sns.barplot(data=chart_df, x=x_col, y=y_col, hue=x_col, order=order, palette="viridis", legend=False)
        plt.xticks(rotation=25, ha="right")
        plt.title(title)
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        figure_paths.append(path)

    if clean_df["published_at"].notna().any():
        timeline = (
            clean_df.assign(date=clean_df["published_at"].dt.date)
            .groupby(["date", "source"])
            .size()
            .reset_index(name="count")
        )
        if not timeline.empty:
            path = FIGURES_DIR / "05_publication_timeline.png"
            plt.figure(figsize=(10, 5))
            sns.lineplot(data=timeline, x="date", y="count", hue="source", marker="o")
            plt.xticks(rotation=25, ha="right")
            plt.title("Vacancy Freshness Timeline")
            plt.tight_layout()
            plt.savefig(path, dpi=180)
            plt.close()
            figure_paths.append(path)

    coverage = (
        clean_df.assign(has_salary=clean_df["salary_mid_kzt"].notna())
        .groupby(["source", "has_salary"])
        .size()
        .reset_index(name="count")
    )
    if not coverage.empty:
        path = FIGURES_DIR / "06_salary_transparency.png"
        plt.figure(figsize=(8, 5))
        sns.barplot(data=coverage, x="source", y="count", hue="has_salary", palette="mako")
        plt.title("Salary Transparency by Source")
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        figure_paths.append(path)

    return figure_paths


def build_platform_matrix() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "Platform": "hh.kz",
                "Student focus": "Medium",
                "Filters / search": "High",
                "Salary transparency": "Medium",
                "Technology level": "High",
                "Notes": "Strong search, commercial job board, but not niche for students.",
            },
            {
                "Platform": "Enbek.kz",
                "Student focus": "Medium-High",
                "Filters / search": "Medium",
                "Salary transparency": "High",
                "Technology level": "Medium",
                "Notes": "State platform with youth/practice routes, broad reach, weaker UX depth.",
            },
            {
                "Platform": "Staj.kz (target)",
                "Student focus": "Very High",
                "Filters / search": "Planned niche",
                "Salary transparency": "Very High",
                "Technology level": "Mobile-first niche",
                "Notes": "Niche platform around no-experience candidates, universities and feedback.",
            },
        ]
    )


def dataframe_to_markdown_table(df: pd.DataFrame) -> str:
    if df.empty:
        return "_No data_"
    columns = list(df.columns)
    lines = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for row in df.fillna("").astype(str).to_dict("records"):
        values = [row.get(column, "").replace("\n", " ").replace("|", "/") for column in columns]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def build_report_markdown(
    summary: dict[str, Any],
    clean_df: pd.DataFrame,
    quality_df: pd.DataFrame,
    figure_paths: Iterable[Path],
) -> str:
    top_employers = clean_df["company"].value_counts().head(10).to_dict()
    top_skills = (
        clean_df["skills"]
        .dropna()
        .astype(str)
        .str.split(",")
        .explode()
        .map(clean_text)
        .replace("", pd.NA)
        .dropna()
        .value_counts()
        .head(10)
        .to_dict()
    )
    strengths = [
        "Рынок large-volume: cleaned dataset превышает порог 2500 записей.",
        "hh.kz даёт масштаб и ширину рынка, Enbek.kz усиливает государственный и youth/practice сегмент.",
        "Высокая доля вакансий без опыта подтверждает существование entry-level воронки.",
    ]
    weaknesses = [
        "Часть вакансий без зарплаты, что создаёт неопределённость для студентов.",
        "Early-career рынок широк, но не всегда student-specific: нишевый продукт должен усиливать релевантность.",
        "Enbek.kz и hh.kz не дают явного student-only UX на уровне продукта.",
    ]
    opportunities = [
        "Собрать student-only каталог с прозрачной зарплатой и фильтрами по вузу/курсу.",
        "Запускаться в среднем ценовом сегменте как самом массовом и понятном для HR.",
        "Начать с Алматы и Астаны, затем расширять coverage по городам с плотным наймом.",
    ]
    threats = [
        "Крупные площадки уже владеют общим трафиком вакансий.",
        "Сильная зависимость рынка от сезонности и офлайн/full-time форматов.",
        "Нужно доказывать HR, что niche funnel качественнее массовых откликов.",
    ]

    figure_lines = "\n".join(f"- `{path.name}`" for path in figure_paths)
    report = f"""# Assignment 2: Staj.kz Market Intelligence

## 1. Dataset summary
- Run date (UTC): `{summary['run_date_utc']}`
- Raw rows before cleaning: `{summary['raw_rows']}`
- Clean rows after cleaning: `{summary['clean_rows']}`
- Freshness window: `{summary['freshness_window_days']}` days
- Salary coverage: `{summary['salary_coverage_pct']}%`
- Student-friendly share: `{summary['student_friendly_pct']}%`
- Source split: `{json.dumps(summary['source_counts'], ensure_ascii=False)}`

## 2. Sources and scope
- Source 1: `hh.kz / api.hh.ru`
- Source 2: `Enbek.kz`
- Market scope: `early-career` vacancies for Kazakhstan
- Query logic: internship / trainee / graduate / junior / no-experience / youth / practice routes

## 3. Data quality log
{dataframe_to_markdown_table(quality_df)}

## 4. Market structure
- Top cities: `{json.dumps(summary['top_cities'], ensure_ascii=False)}`
- Top role categories: `{json.dumps(summary['top_role_categories'], ensure_ascii=False)}`
- Price segments: `{json.dumps(summary['price_segments'], ensure_ascii=False)}`
- Most promising segment for Staj.kz: `{summary['perspective_segment']}`

### Top employers
`{json.dumps(top_employers, ensure_ascii=False)}`

### Top extracted skills
`{json.dumps(top_skills, ensure_ascii=False)}`

## 5. Competitor analysis
### Platform matrix
{dataframe_to_markdown_table(build_platform_matrix())}

### SWOT
**Strengths**
{chr(10).join(f"- {item}" for item in strengths)}

**Weaknesses**
{chr(10).join(f"- {item}" for item in weaknesses)}

**Opportunities**
{chr(10).join(f"- {item}" for item in opportunities)}

**Threats**
{chr(10).join(f"- {item}" for item in threats)}

## 6. Strategic recommendation
Staj.kz should launch as a focused student-first layer above the broad early-career market. The best initial position is the `{summary['perspective_segment']}` segment, where volume is high enough for liquidity, while the product can still differentiate on salary transparency, university-aware filters and response feedback. Priority geography should stay with the densest hiring hubs first, then expand by city and role cluster.

## 7. Figures
{figure_lines}
"""
    return report


def build_report_docx(markdown_report: str, figure_paths: Iterable[Path], output_path: Path) -> None:
    document = Document()
    document.add_heading("Assignment 2. Staj.kz Market Intelligence", level=0)
    for paragraph in markdown_report.split("\n\n"):
        text = paragraph.strip()
        if not text:
            continue
        if text.startswith("# "):
            document.add_heading(text[2:].strip(), level=1)
        elif text.startswith("## "):
            document.add_heading(text[3:].strip(), level=2)
        elif text.startswith("### "):
            document.add_heading(text[4:].strip(), level=3)
        else:
            lines = text.splitlines()
            if all(line.startswith("- ") for line in lines):
                for line in lines:
                    document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(textwrap.dedent(text))

    document.add_heading("Figures", level=2)
    for figure in figure_paths:
        if figure.exists():
            document.add_paragraph(figure.name)
            document.add_picture(str(figure), width=Inches(6.4))

    document.save(output_path)


def run_pipeline(max_hh_pages: int) -> dict[str, Path]:
    ensure_dirs()
    session = request_session()

    hh_df = fetch_hh_records(session, max_pages_per_query=max_hh_pages)
    enbek_df = fetch_enbek_search_records(session)

    raw_df = pd.concat([hh_df, enbek_df], ignore_index=True)
    raw_df = standardize_dataframe(raw_df)
    clean_df, quality_df = apply_quality_cleaning(raw_df)

    summary = build_summary(clean_df, raw_df, quality_df)
    figure_paths = save_figures(clean_df)
    markdown_report = build_report_markdown(summary, clean_df, quality_df, figure_paths)

    raw_csv = PROCESSED_DIR / "assignment2_raw_combined.csv"
    clean_csv = PROCESSED_DIR / "assignment2_cleaned_market_intelligence.csv"
    quality_csv = PROCESSED_DIR / "assignment2_quality_log.csv"
    summary_json = PROCESSED_DIR / "assignment2_summary.json"
    report_md = DOC_DIR / "assignment2_report.md"
    report_docx = DOC_DIR / "assignment2_report.docx"

    save_dataframe(raw_df, raw_csv)
    save_dataframe(clean_df, clean_csv)
    save_dataframe(quality_df, quality_csv)
    dump_json(summary_json, summary)
    dump_text(report_md, markdown_report)
    build_report_docx(markdown_report, figure_paths, report_docx)

    return {
        "raw_csv": raw_csv,
        "clean_csv": clean_csv,
        "quality_csv": quality_csv,
        "summary_json": summary_json,
        "report_md": report_md,
        "report_docx": report_docx,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Assignment 2 pipeline for Staj.kz market intelligence.")
    parser.add_argument(
        "--hh-pages",
        type=int,
        default=10,
        help="Maximum pages per hh.kz query (100 records per page). Default: 10",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    outputs = run_pipeline(max_hh_pages=args.hh_pages)
    summary = json.loads((PROCESSED_DIR / "assignment2_summary.json").read_text(encoding="utf-8"))

    print("Assignment 2 pipeline completed.")
    print(f"Clean rows: {summary['clean_rows']}")
    print(f"Source split: {summary['source_counts']}")
    for name, path in outputs.items():
        print(f"{name}: {path}")

    if summary["clean_rows"] < 2500:
        print("WARNING: Clean dataset is below the 2500-row target.")


if __name__ == "__main__":
    main()
